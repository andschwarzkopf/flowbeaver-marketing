#!/usr/bin/env python3
"""Build print-ready DOCX from the AVV markdown template, applying Flowbeaver branding."""

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(__file__).resolve().parent.parent / "rechtliches" / "AVV-Template-Flowbeaver.md"
DST = Path(__file__).resolve().parent.parent / "rechtliches" / "AVV-Template-Flowbeaver.docx"

# Flowbeaver brand colors
PURPLE = RGBColor(0x52, 0x26, 0x7A)
DARK_RED = RGBColor(0x3C, 0x0F, 0x04)
BLACK = RGBColor(0x00, 0x00, 0x00)
TEXT = RGBColor(0x40, 0x40, 0x40)
BORDER = RGBColor(0xE0, 0xE0, 0xE0)
WARN_BG = "F4EEF8"  # very light purple tint for the disclaimer box

FONT = "DM Sans"
FONT_FALLBACK = "Calibri"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex="E0E0E0", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_horizontal_rule(paragraph, color_hex="E0E0E0"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_run(run, *, size=11, bold=False, color=TEXT, font=FONT):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE = re.compile(r"`([^`]+)`")


def add_inline(paragraph, text, *, size=11, color=TEXT, base_bold=False):
    """Render a markdown line with **bold** and `code` segments to a paragraph."""
    if not text:
        return
    # First split on bold, then within each non-bold chunk split on code.
    pos = 0
    segments = []
    for m in INLINE_BOLD.finditer(text):
        if m.start() > pos:
            segments.append((text[pos : m.start()], False))
        segments.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False))

    for seg_text, is_bold in segments:
        # split out inline `code` (rendered as monospace, no special bg)
        sub_pos = 0
        for cm in INLINE_CODE.finditer(seg_text):
            if cm.start() > sub_pos:
                run = paragraph.add_run(seg_text[sub_pos : cm.start()])
                style_run(run, size=size, bold=base_bold or is_bold, color=color)
            run = paragraph.add_run(cm.group(1))
            style_run(run, size=size - 1, bold=base_bold or is_bold, color=color, font="Consolas")
            sub_pos = cm.end()
        if sub_pos < len(seg_text):
            run = paragraph.add_run(seg_text[sub_pos:])
            style_run(run, size=size, bold=base_bold or is_bold, color=color)


def setup_document():
    doc = Document()

    # Page setup: A4, 2.5 cm margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), FONT)

    # Footer with page numbers
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Flowbeaver — AVV Vorlage  ·  Seite ")
    style_run(run, size=9, color=TEXT)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = fp.add_run()
    style_run(page_run, size=9, color=TEXT)
    page_run._element.append(fld_begin)
    page_run._element.append(instr)
    page_run._element.append(fld_end)
    run = fp.add_run(" von ")
    style_run(run, size=9, color=TEXT)
    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.text = "NUMPAGES"
    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    pages_run = fp.add_run()
    style_run(pages_run, size=9, color=TEXT)
    pages_run._element.append(fld_begin2)
    pages_run._element.append(instr2)
    pages_run._element.append(fld_end2)

    return doc


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=24, bold=True, color=DARK_RED)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(20)
    add_inline(p, text, size=14, color=PURPLE, base_bold=True)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=15, bold=True, color=PURPLE)
    add_horizontal_rule(p, "52267A")


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=12, bold=True, color=BLACK)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=11, bold=True, color=BLACK)


def add_para(doc, text, *, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    # Preserve internal line breaks (e.g. multi-line address blocks) as soft breaks.
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            br_run = p.add_run()
            br = OxmlElement("w:br")
            br_run._element.append(br)
            style_run(br_run, size=size)
        add_inline(p, part, size=size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    # List Bullet style already adds a bullet — but ensure font/color
    add_inline(p, text, size=11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    add_inline(p, text, size=11)
    return p


def add_disclaimer_box(doc, text):
    """One-cell table with light purple background, used for the legal disclaimer."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, WARN_BG)
    set_cell_borders(cell, "52267A", size="6")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_inline(p, text, size=10, color=DARK_RED)
    # spacing after the box
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_signature_block(doc, lines):
    """Two-column signature table at the bottom of the contract."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    add_inline(p, lines[0], size=11, base_bold=True)  # Ort, Datum

    table = doc.add_table(rows=3, cols=2)
    table.autofit = True
    for col, label in enumerate(lines[1:]):
        # signature line
        sig_cell = table.cell(0, col)
        sp = sig_cell.paragraphs[0]
        sp.paragraph_format.space_before = Pt(40)
        run = sp.add_run("_" * 40)
        style_run(run, size=11, color=BLACK)
        # role label
        role_cell = table.cell(1, col)
        rp = role_cell.paragraphs[0]
        add_inline(rp, label["role"], size=10, base_bold=True)
        # name line
        name_cell = table.cell(2, col)
        np_ = name_cell.paragraphs[0]
        add_inline(np_, label["name"], size=10)
    # remove borders for signature table
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tc_borders.append(b)
            tc_pr.append(tc_borders)


def add_table_from_md(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.autofit = True
    # header
    for i, h in enumerate(header):
        cell = table.cell(0, i)
        set_cell_shading(cell, "52267A")
        set_cell_borders(cell, "52267A", size="6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h.strip())
        style_run(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # data rows
    for r_i, row in enumerate(rows, start=1):
        for c_i, cell_text in enumerate(row):
            cell = table.cell(r_i, c_i)
            set_cell_borders(cell, "E0E0E0", size="4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, cell_text.strip(), size=10)
    # spacing after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


# ---- Markdown parser (purpose-built for this document) ----

def parse_md(md_text):
    """Yield blocks: ('title', text), ('subtitle', text), ('h1', text), ('h2', text),
    ('h3', text), ('para', text), ('bullet', text), ('numbered', text),
    ('blockquote', text), ('hr',), ('pagebreak',), ('signature', dict),
    ('table', header, rows)."""
    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    def is_table_row(s):
        return s.strip().startswith("|") and s.strip().endswith("|")

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank
        if not stripped:
            i += 1
            continue

        # Horizontal rule (---). Two HRs in a row → page break (per markdown convention used in source).
        if stripped == "---":
            # peek for double HR (`---\n---\n`)
            j = i + 1
            while j < n and not lines[j].strip():
                j += 1
            if j < n and lines[j].strip() == "---":
                yield ("pagebreak",)
                i = j + 1
                continue
            yield ("hr",)
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            yield ("title", stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("## "):
            yield ("h1", stripped[3:].strip())
            i += 1
            continue
        if stripped.startswith("### "):
            yield ("h2", stripped[4:].strip())
            i += 1
            continue
        if stripped.startswith("#### "):
            yield ("h3", stripped[5:].strip())
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            yield ("blockquote", " ".join(buf).strip())
            continue

        # Bullet list
        if stripped.startswith("- "):
            yield ("bullet", stripped[2:].strip())
            i += 1
            continue

        # Numbered list (1.  2.  …)
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            yield ("numbered", m.group(1))
            i += 1
            continue

        # Tables
        if is_table_row(stripped):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 1
            # skip separator row
            if i < n and re.match(r"^\|[\s:|-]+\|$", lines[i].strip()):
                i += 1
            rows = []
            while i < n and is_table_row(lines[i].strip()):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            yield ("table", header, rows)
            continue

        # Paragraph: collect consecutive non-blank lines
        buf = [stripped]
        i += 1
        while i < n:
            s = lines[i].strip()
            if not s:
                break
            if s.startswith(("# ", "## ", "### ", "#### ", "- ", ">", "|")) or s == "---":
                break
            if re.match(r"^\d+\.\s+", s):
                break
            buf.append(s)
            i += 1
        yield ("para", "\n".join(buf))


def render(doc, blocks):
    """Render the parsed blocks. Custom handling for the Flowbeaver disclaimer
    and signature block."""
    blocks = list(blocks)
    idx = 0

    while idx < len(blocks):
        block = blocks[idx]
        kind = block[0]

        # First H1 (Vertrag zur Auftragsverarbeitung) → title
        if kind == "title":
            add_title(doc, block[1])
            idx += 1
            # Look ahead: next 'para' starting with **gemäß is the subtitle
            while idx < len(blocks) and blocks[idx][0] in ("hr",):
                idx += 1
            if idx < len(blocks) and blocks[idx][0] == "para" and "DSGVO" in blocks[idx][1]:
                add_subtitle(doc, blocks[idx][1])
                idx += 1
            continue

        if kind == "blockquote":
            add_disclaimer_box(doc, block[1])
            idx += 1
            continue

        if kind == "hr":
            # Soft visual divider
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run("◆ ◆ ◆")
            style_run(run, size=9, color=PURPLE)
            idx += 1
            continue

        if kind == "pagebreak":
            doc.add_page_break()
            idx += 1
            continue

        if kind == "h1":
            add_h1(doc, block[1])
            idx += 1
            continue

        if kind == "h2":
            add_h2(doc, block[1])
            idx += 1
            continue

        if kind == "h3":
            add_h3(doc, block[1])
            idx += 1
            continue

        if kind == "bullet":
            add_bullet(doc, block[1])
            idx += 1
            continue

        if kind == "numbered":
            add_numbered(doc, block[1])
            idx += 1
            continue

        if kind == "table":
            _, header, rows = block
            add_table_from_md(doc, header, rows)
            idx += 1
            continue

        if kind == "para":
            text = block[1]
            # Detect signature block: "[Ort], [Datum]" pattern, followed by underline lines
            if text.startswith("**[Ort]") or text == "**[Ort], [Datum]**":
                # Render as small signature block manually using the next few lines
                add_signature_block(
                    doc,
                    [
                        text,
                        {"role": "Für den Verantwortlichen", "name": "[Name, Funktion]"},
                        {"role": "Für den Auftragsverarbeiter", "name": "Andreas Schwarzkopf, Flowbeaver"},
                    ],
                )
                # skip the underline-paragraph block(s) that follow in source
                idx += 1
                while idx < len(blocks) and blocks[idx][0] == "para" and (
                    "_" in blocks[idx][1] or "Für den" in blocks[idx][1] or "Schwarzkopf" in blocks[idx][1]
                ):
                    idx += 1
                continue

            # Skip leftover underline-only paragraphs from source
            if re.match(r"^_{5,}", text):
                idx += 1
                continue

            # "Ende des Dokuments" → centered, italic
            if "Ende des Dokuments" in text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(20)
                run = p.add_run("Ende des Dokuments")
                style_run(run, size=10, color=PURPLE, bold=True)
                idx += 1
                continue

            add_para(doc, text)
            idx += 1
            continue

        idx += 1


def main():
    md_text = SRC.read_text(encoding="utf-8")
    doc = setup_document()
    blocks = parse_md(md_text)
    render(doc, blocks)
    doc.save(DST)
    print(f"Wrote {DST} ({DST.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
